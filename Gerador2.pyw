import sys
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import (
    QApplication,
    QWidget,
    QLabel,
    QLineEdit,
    QPushButton,
    QVBoxLayout,
    QMessageBox,
)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from datetime import datetime
import os

class GeradorDeclaracao(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Gerador de Declaração")
        self.setWindowIcon(QIcon('C:/Users/danie/Downloads/jornada-de-trabalho.png'))
        self.setStyleSheet("background-color: black; color: white;")
        layout = QVBoxLayout()

        self.entry_nome = QLineEdit()
        self.entry_endereco = QLineEdit()
        self.entry_cpf = QLineEdit()
        self.entry_cep = QLineEdit()
        
        self.entry_nome.setStyleSheet("color: magenta;")
        self.entry_endereco.setStyleSheet("color: magenta;")
        self.entry_cpf.setStyleSheet("color: magenta;")
        self.entry_cep.setStyleSheet("color: magenta;")


        layout.addWidget(QLabel("Nome:"))
        layout.addWidget(self.entry_nome)
        layout.addWidget(QLabel("Endereço:"))
        layout.addWidget(self.entry_endereco)
        layout.addWidget(QLabel("CPF:"))
        layout.addWidget(self.entry_cpf)
        layout.addWidget(QLabel("CEP:"))
        layout.addWidget(self.entry_cep)

        btn_gerar = QPushButton("Gerar Documento")
        btn_gerar.clicked.connect(self.gerar_documento)
        btn_gerar.setStyleSheet("background-color: white; color: black; font-weight: bold; font-family: Arial;")
        layout.addWidget(btn_gerar)

        btn_cancelar = QPushButton("Cancelar")
        btn_cancelar.clicked.connect(self.cancelar)
        btn_cancelar.setStyleSheet("background-color: red; color: white; font-weight: bold; font-family: Arial;")
        layout.addWidget(btn_cancelar)
        

        self.setLayout(layout)
        self.setGeometry(900, 300, 270, 300)

    def gerar_documento(self):
        nome = self.entry_nome.text()
        endereco = self.entry_endereco.text()
        cpf = self.entry_cpf.text()
        cep = self.entry_cep.text()

        try:
            cpf = int(cpf)
            cpf_string = str(cpf)
            cpf_formatado = f"{cpf_string[:3]}.{cpf_string[3:6]}.{cpf_string[6:9]}-{cpf_string[9:]}"

            cep = int(cep)
            cep_string = str(cep)
            cep_formatado = f"{cep_string[:5]}-{cep_string[5:]}"
        except ValueError:
            QMessageBox.critical(self, "Erro", "CPF e CEP devem conter apenas números")
            return

        data_atual = datetime.now()
        dia = data_atual.day
        mes = data_atual.strftime("%B")
        ano = data_atual.year

        doc = Document()
        section = doc.sections[0]
        section.page_width = Pt(21 * 28.35)   
        section.page_height = Pt(29.7 * 28.35)  
        section.top_margin = Cm(2.5)  
        section.bottom_margin = Cm(2.5)  
        section.left_margin = Cm(3)  
        section.right_margin = Cm(3)  

        texto1 = '''Associação de Moradores da Comunidade dos Teixeira e Adjacências Sede: Estrada dos Teixeira nº: 1200 lote: 200 cep: 22723-205 Taquara-RJ CNPJ: 51.813.448/0001-65 tel:(21)988251763'''

        texto2 = 'Declaração'
        texto3 = f'''Declaro para os devidos fins e fazer prova junto aos
Órgãos Públicos Federais, Estaduais e Municipais, e junto
também às Instituições financeiras, bancárias e privadas que, {nome} Registro geral- CPF:
{cpf_formatado} Detran RJ, reside na Estrada dos Teixeira {endereco} CEP: {cep_formatado} - Taquara - Jacarepaguá - Rio de Janeiro.'''
        texto4 = f''' _____________________________________________  
Hélio Monteiro da Silva'''
        texto5 = f'''
Presidente da Associação de Moradores da Comunidade dos Teixeira e Adjacências.
Rio de Janeiro, {dia} de {mes} de {ano}.'''

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
        run = paragraph.add_run(texto1)
        run.bold = False
        run.font.size = Pt(14)
        run.font.name = 'Calibri'
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(12)

        doc.add_paragraph('''
        ''')

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
        run = paragraph.add_run(texto2)
        run.bold = True
        run.font.size = Pt(17)
        run.font.name = 'Calibri'
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(12)

        doc.add_paragraph('''
        ''')

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
        run = paragraph.add_run(texto3)
        run.bold = False
        run.font.size = Pt(16)
        run.font.name = 'Calibri'
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(15)
        doc.add_paragraph('''
                        
                        
                        ''')

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
        run = paragraph.add_run(texto4)
        run.bold = False
        run.font.size = Pt(16)
        run.font.name = 'Calibri'
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(12)

        doc.add_paragraph('''
                        
                        
                        
                        
                        
                        ''')

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
        run = paragraph.add_run(texto5)
        run.bold = False
        run.font.size = Pt(16)
        run.font.name = 'Calibri'
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(12)

        # Diretório onde o arquivo será salvo
        diretorio_salvamento = "C:/Users/danie/Desktop/Declaracções geradas"

        # Verifica se o diretório existe, se não existir, cria
        if not os.path.exists(diretorio_salvamento):
            os.makedirs(diretorio_salvamento)

        # Define o nome completo do arquivo, incluindo o caminho do diretório
        nome_arquivo = os.path.join(diretorio_salvamento, f"documento_{nome}.docx")

        # Salva o arquivo no diretório especificado
        doc.save(nome_arquivo)
        QMessageBox.information(self, "Sucesso", f"Arquivo DOCX '{nome_arquivo}' criado com sucesso!")

    def cancelar(self):
        self.close()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    janela = GeradorDeclaracao()
    janela.show()
    sys.exit(app.exec_())
